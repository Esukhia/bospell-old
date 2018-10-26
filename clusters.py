import time
import os
import re

import hashids

import threading
import multiprocess

from pybo import *

from sympound import sympound
from jellyfish import levenshtein_distance

import docx

#import ngram
#import pickle

debug = False

class Symspell:
    def __init__(self, distancefun=levenshtein_distance, mded=5):
        self.ssc = sympound(distancefun=distancefun, maxDictionaryEditDistance=mded)

        """self.symspell_thread = threading.Thread(target=self.load_dict)
        self.symspell_thread.start()"""
        self.load_dict

    def load_dict(self):
        if not os.path.isfile("symspell.pickle"):
            self.ssc.load_dictionary(os.path.join("resources", "symspell", "gmd.txt"))
            self.ssc.save_pickle("symspell.pickle")
        else:
            self.ssc.load_pickle("symspell.pickle")

    def comp(self, string):
        #self.symspell_thread.join()
        result = self._symspell_lookup(string)
        return True if result == string else result

    def _symspell_lookup(self, string, edm=4):
        r = self.ssc.lookup_compound(input_string=string, edit_distance_max=edm)
        return r.term


class Matcher:
    """Class Matcher to execute all the tests.

    The profile POS is used by default by the pybo Tokenizer

    Args:
        profile (string): name of the profile used by pybo

    Attributes:
        FML (int): limit between the low and the high frequency of a word
        tok (list of :obj:`tokens`): the tokens returned by BoTokenizer within pybo module
        rule_list (:dict:`string` of tupple:int, string, int, string): used to correct possible recognition errors of namsel-ocr
            key as `string`: regex of the search syllable
            tupple:
                `int` == -1: depends on the last syllable
                        == 0: depends on itself
                        == 1: depends on the next syllable
                ``string`: regex of the dependent syllable
                `int` == 0 the found syllable should be replaced
                        == 1 the next syllable should be replaced
                `string`: the replace value
    """
    def __init__(self, profile="MGD"):
        self.SKRT_CONS = 5
        self.SKRT_SUB_CONS = 6
        self.SKRT_VOW = 7

        self.NON_WORD = 'non-word'
        self.FML = 49

        self.init_tokenizer(profile)

        self.rule_list = []
        self.error_list = []
        self.warning_list = []
        self.__rule_list()

        self.format_style = {"auto-correct":(178, 178, 178), "error":(255, 0, 0), "warning":(255, 153, 204),
                                        "red":(255, 0, 0), "green":(0, 255, 0), "blue":(0, 0, 255), "orange":(255, 192, 0),
                                        "purple":(204, 0, 204)}

    def __rule_list(self):
        with open("rule_list.txt", mode="r", encoding="utf-8") as f:
            for line in f.read().split('\n'):
                if not line or line.startswith("#") : continue

                if line.startswith("@"):
                    self.error_list.append(list(line[1:].split(",")))
                elif line.startswith("~"):
                    self.warning_list.append(list(line[1:].split(",")))
                else:
                    s, c = line.split("->")
                    self.rule_list.append([list(s.split(",")), list(c.split(","))])

    def init_tokenizer(self, profile):
        self.tok = BoTokenizer(profile)

    @staticmethod
    def is_punct_token(token):
        return token.tag == 'punct'

    @staticmethod
    def is_monosyl_token(token):
        return token.syls and len(token.syls) == 1 and token.pos != "PART"

    def is_non_word_token(self, token):
        return token.tag == self.NON_WORD

    def is_non_bo_or_number(self, token):
        return token.tag == "non-bo" or token.tag == "num"

    @staticmethod
    def is_oov_token(token):
        return token.tag == "oov"

    def is_mono_or_punct_or_nonword(self, token):
        return self.is_monosyl_token(token) or \
               self.is_punct_token(token)

    def is_HFM(self, token):
        return token.freq > self.FML

    def is_LFM(self, token):
        return token.freq <= self.FML

    @staticmethod
    def has_skrt_token_att(token):
        return token.skrt

    @staticmethod
    def cluster_test(i, idx):
        # return 2 for pop and underline and 1 for underline only
        return 2 if idx[1] == 1 and i == idx[0] \
            else 1 if idx[1] == 2 and i in range(idx[0][0], idx[0][1]) \
            else 2 if idx[1] == 2 and i == idx[0][1] \
            else 0

class Structure:
    def __init__(self, page):
        self.page = page
        self.hash = hashids.Hashids(salt=hex(id(self)), min_length=9)
        self.hash_cluster_list = []
        self.clusters = []

        self.tokens = matcher.tok.tokenize(self.page)

        self.clusterize()
        self.adjust_clusters()
        self.cluster_format()

    def clusterize(self):
        tmp = []
        for i, token in enumerate(self.tokens):
            if matcher.is_mono_or_punct_or_nonword(token) and \
                    not (not tmp and matcher.is_punct_token(token)):
                tmp.append(i)
                if i == len(self.tokens) - 1:
                    self.clusters.append([tmp[0], i])
            elif tmp:
                self.clusters.append([tmp[0], tmp[-1]])
                tmp = []
        if debug:
            print(self.clusters)

    def get_token(self, index):
        return self.tokens[index].content

    def adjust_clusters(self):
        for i, cluster in enumerate(self.clusters):
            left_cluster_id = cluster[0]
            right_cluster_id = cluster[1]

            # Left context
            while left_cluster_id < right_cluster_id:
                if self.get_mono_token_freq(left_cluster_id) > matcher.FML:
                    left_cluster_id += 1
                    while left_cluster_id < len(self.tokens) and matcher.is_punct_token(self.tokens[left_cluster_id]):
                        left_cluster_id += 1
                else: break

            if left_cluster_id == right_cluster_id:
                if self.get_mono_token_freq(left_cluster_id) > matcher.FML:
                    self.clusters[i] = None
                else:
                    self.clusters[i] = [left_cluster_id]
                continue
            elif left_cluster_id > right_cluster_id:
                self.clusters[i] = None
                continue
            else:
                self.clusters[i][0] = left_cluster_id

            # Right context
            while right_cluster_id > left_cluster_id:
                while matcher.is_punct_token(self.tokens[right_cluster_id]):
                    right_cluster_id -= 1
                if self.get_mono_token_freq(right_cluster_id) > matcher.FML:
                    right_cluster_id -= 1
                else: break

            if right_cluster_id == left_cluster_id:
                if self.get_mono_token_freq(right_cluster_id) > matcher.FML:
                    self.clusters[i] = None
                else:
                    self.clusters[i] =  [right_cluster_id]
                continue
            elif right_cluster_id < left_cluster_id:
                self.clusters[i] = None
                continue
            else:
                self.clusters[i][1] = right_cluster_id

        self.clusters = [c for c in self.clusters if c is not None]

    def get_mono_token_freq(self, index):
        ti = self.tokens[index]
        if matcher.is_monosyl_token(ti):
            if hasattr(ti, "freq") and ti.freq:
                return ti.freq
            else: return 0

    def cluster_format(self):
        for i, cluster_index in enumerate(self.clusters):
            type = self.cluster_type(cluster_index)
            self.hash_cluster_list.append(self.hash.encode(cluster_index[0]))
            if len(cluster_index) == 1:
                self.clusters[i] = (cluster_index[0], 1, type)
            else:
                self.clusters[i] = (tuple(cluster_index), 2, type)

    def cluster_type(self, cluster_index):
        type = []
        if len(cluster_index) > 1:
            cluster_token = self.tokens[cluster_index[0]:cluster_index[1]+1]
        else:
            cluster_token = [self.tokens[cluster_index[0]]]

        for t in cluster_token:
            if matcher.is_non_word_token(t):
                type.append('non-word')
            elif matcher.has_skrt_token_att(t):
                type.append("skrt")
            elif matcher.is_oov_token(t):
                type.append("oov")
            elif matcher.is_monosyl_token(t):
                type.append("mono")

        if any([t == 'non-word' for t in type]): return 'non-word'
        elif all([t == "skrt" for t in type]): return "skrt"
        elif any([t == "skrt" for t in type]): return "both"
        elif any([t == 'oov' for t in type]): return 'oov'
        elif all([t == "mono"for t in type]): return "mono"
        else: return "other"

    def cql_test(self, list_, code, correcting=False):
        for q in list_:
            ending = "་?"
            if correcting:
                q, c = q
            if q[0].startswith("°"):
                q[0] = q[0][1:]
                ending = ""
            matcher1 = CQLMatcher("".join(['[content="' + q[i] + ending +'"] ' for i in range(len(q))]))
            slices = matcher1.match(self.tokens)
            if slices:
                for start, end in slices:
                    for idx, i in enumerate(range(start, end + 1)):
                        self.tokens[i]._["format"] = code
                        if correcting and c[idx]:
                            self.tokens[i].content = "" if c[idx] == "^" else re.sub(r"" + q[idx], r"" + c[idx], self.tokens[i].content)
                            if debug:
                                print(self.tokens[i].content)

    def highlight(self, destination="html"):
        if self.clusters:
            clust = list(reversed(self.clusters))
            for i, t in enumerate(self.tokens):
                if len(clust) and (i == 0 or test == 2):
                    cluster_index = clust.pop()
                test = matcher.cluster_test(i, cluster_index)
                if test:
                    if cluster_index[2] == "non-word": t._["format"] = "red"
                    elif cluster_index[2] == "oov": t._["format"] = "orange"
                    elif cluster_index[2] == "skrt": t._["format"] = "blue"
                    elif cluster_index[2] == "mono": t._["format"] = "purple"
                    elif cluster_index[2] == "both": t._["format"] = "green"
                else: t._["format"] = ""

        self.cql_test(matcher.error_list, "error")
        self.cql_test(matcher.warning_list, "warning")
        self.cql_test(matcher.rule_list, "auto-correct", True)

        self.add_to_docx() if destination == "docx" else self.add_to_html()

    def add_to_docx(self):
        # Save to Docx
        docx_result = docx.Document()
        p = docx_result.add_paragraph("")

        for i, t in enumerate(self.tokens):
            f = t._.get("format")
            wp = p.add_run(t.content)
            if matcher.is_non_bo_or_number(t):
                wp.font.color.rgb = docx.shared.RGBColor(0, 255, 255)
            elif not f:
                wp.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            else:
                for k, fs in matcher.format_style.items():
                    if f == k:
                        r, g, b = fs
                        wp.font.color.rgb = docx.shared.RGBColor(r, g, b)
                        break
            wp.font.size = docx.shared.Pt(22)
        docx_result.save('result.docx')

    def add_to_html(self):
        # Save to html
        html_result = ["<p style='font-size:22px'>"]
        for i, t in enumerate(self.tokens):
            f = t._.get("format")
            if matcher.is_non_bo_or_number(t):
                html_result.append(f"<span style='color:rgb(0, 255, 255);'>{t.content}</span>")
            elif not f:
                html_result.append(t.content)
            else:
                html_result.append(f"<span style='color:rgb{matcher.format_style[f]};'>{t.content}</span>")
        html_result.append("</p>")
        print("".join(html_result))


class Prediction:
    def __init__(self, clusters, tokens):
        self.clusters = clusters
        self.tokens = tokens
        self.limit = -1
        self.result = {}
        self.rr = []

    def adjust_clusters(self):
        self.token_list = []
        for c in self.clusters:
            if isinstance(c[0], tuple):
                self.left = c[0][0]
                self.right = c[0][1]
            else:
                self.left = self.right = c[0]

                self.left = self.get_previous_token(self.left)
                self.right = self.get_next_token(self.right)

            while True:
                self.cluster_tokens = self.load_tokens()
                if self.cluster_tokens == -1: break
                self.token_list.append(self.cluster_tokens)

    def add_saved_tocken_info(self, tmp):
        if tmp:
            self.add_left_token = tmp[0]
            self.add_right_token = tmp[-1]

    def load_tokens(self, last_token="", limit=3):
        if not last_token:
                if self.right - self.left > limit:
                    for i, c in enumerate(self.tokens[self.left:self.right+1][0:limit]):
                        if matcher.is_punct_token(c):
                            self.limit = i
                            tmp = self.tokens[self.left:self.right+1][0:i]
                            self.left = self.limit+1
                            while matcher.is_punct_token(self.tokens[self.left]):
                                self.left += 1
                            self.add_saved_tocken_info(tmp)
                            return tmp
                    else:
                        tmp = self.tokens[self.left:self.right+1][0:limit]
                        self.left += limit
                        self.add_saved_tocken_info(tmp)
                        return tmp
                elif self.right - self.left < 0:
                    return -1
                else:
                    tmp = self.tokens[self.left:self.right+1]
                    self.left = self.right + 1
                    self.add_saved_tocken_info(tmp)
                    return tmp

    def get_previous_token(self, index):
        if index > 0 and not matcher.is_punct_token(self.tokens[index-1]):
            return index-1
        else: return index

    def get_next_token(self, index):
        if index < len(self.tokens)-1 and not matcher.is_punct_token(self.tokens[index+1]):
            return index+1
        else: return index

def prediction():
    sy = Symspell()
    def multiprocess_func(x):
        return sy.comp(x)

    p = Prediction(s.clusters, s.tokens)
    p.adjust_clusters()

    test = []
    #hashing = []
    for tl in p.token_list:
        test.append("".join([token.cleaned_content for token in tl]))
    #hashing.append("".join([hash for hash in s.hash_cluster_list]))

    print(test)
    #print(hashing)

    with multiprocess.Pool() as p:
        results = p.map(multiprocess_func, test)
        print(results)
    #print(hashing)


if __name__ == '__main__':
    start_time = time.time()

    matcher = Matcher()

    doc = docx.Document('demo.docx')

    page = "\n".join([tibetan.text for tibetan in doc.paragraphs])

    if debug:
        print(f"\n{page}")

    s = Structure(page)
    if debug:
        for t in s.tokens:
            print(f"TOKEN: {s.tokens.index(t)} content: {t.content} /{t.phono}/ tag: {t.tag} skrt: {t.skrt} freq: {t.freq} pos: {t.pos}")
        print(f"\n{s.clusters}")

    s.highlight("docx")

    #prediction()

    print('%2.2f sec' % (time.time() - start_time))